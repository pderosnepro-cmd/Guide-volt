#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour générer un document Word professionnel à partir du fichier Excel
"Données de l'entreprise (12).xlsx".

Structure du document :
- Chaque onglet Excel devient un thème dans le Word.
- Chaque ligne devient une sous-section avec ses caractéristiques.
- Le mode (Avoir/Compensation/Normal) est déterminé automatiquement.
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

# ============================================================================
# CONFIGURATION
# ============================================================================

# Chemin du fichier Excel
EXCEL_FILE = "Données de l'entreprise (12).xlsx"

# Chemin du fichier Word de sortie
OUTPUT_WORD_FILE = f"Configuration_Zeus_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

# Couleurs pour les styles
PRIMARY_COLOR = RGBColor(0, 51, 102)  # Bleu foncé
SECONDARY_COLOR = RGBColor(102, 102, 102)  # Gris

# ============================================================================
# FONCTIONS POUR DÉTERMINER LE MODE
# ============================================================================

def determine_mode(row, sheet_name, contrats_evenements_df=None, contrats_variables_df=None, compteurs_speciaux_df=None):
    """
    Détermine le mode (Avoir/Compensation/Normal) selon les règles :
    - Avoir si règle d'acquisition présente
    - Compensation si lié à un compteur spécial
    - Normal sinon
    """
    # Cas 1 : Événements journaliers
    if sheet_name == "Événements journaliers":
        if pd.notna(row.get("Règle d'acquisition", None)) and str(row["Règle d'acquisition"]).strip() != "":
            return "Avoir"
        else:
            # Vérifier si lié à un compteur spécial via les contrats
            if contrats_evenements_df is not None and not contrats_evenements_df.empty:
                event_id = row.get("Identifiant")
                if event_id in contrats_evenements_df["Identifiant"].values:
                    # Vérifier si le compteur à déduire est un compteur spécial
                    compteur_a_deduire = contrats_evenements_df[
                        contrats_evenements_df["Identifiant"] == event_id
                    ]["Compteur à déduire"].iloc[0]
                    if pd.notna(compteur_a_deduire) and str(compteur_a_deduire).strip() != "":
                        return "Compensation"
            return "Normal"
    
    # Cas 2 : Liste des compteurs spéciaux
    elif sheet_name == "Liste des compteurs spéciaux":
        return "Compensation"
    
    # Cas 3 : Contrats - événements
    elif sheet_name == "Contrats - événements":
        if pd.notna(row.get("Règle d'acquisition", None)) and str(row["Règle d'acquisition"]).strip() != "":
            return "Avoir"
        else:
            compteur_a_deduire = row.get("Compteur à déduire")
            if pd.notna(compteur_a_deduire) and str(compteur_a_deduire).strip() != "":
                return "Compensation"
            return "Normal"
    
    # Cas 4 : Contrats - variables de paie
    elif sheet_name == "Contrats - variables de paie":
        return "Compensation"  # Les variables de paie sont souvent liées à des compteurs
    
    # Cas 5 : MàJ compte événements/variables
    elif sheet_name in ["MàJ compte événements", "MàJ compte variables de paie"]:
        # Ces onglets ne définissent pas de mode, on retourne None
        return None
    
    # Cas par défaut
    return "Normal"


# ============================================================================
# FONCTIONS POUR EXTRAIRE LES DONNÉES
# ============================================================================

def load_excel_data(file_path):
    """Charge tous les onglets du fichier Excel dans un dictionnaire de DataFrames."""
    xls = pd.ExcelFile(file_path)
    sheets = {sheet_name: xls.parse(sheet_name) for sheet_name in xls.sheet_names}
    return sheets


def get_export_paie_mapping(sheets):
    """Crée un dictionnaire pour mapper les identifiants aux codes export paie."""
    export_paie_evenements = sheets.get("Export paie - événements", pd.DataFrame())
    export_paie_variables = sheets.get("Export paie - variables de paie", pd.DataFrame())
    
    # Dictionnaire pour les événements
    export_paie_evenements_dict = {}
    if not export_paie_evenements.empty:
        for _, row in export_paie_evenements.iterrows():
            export_paie_evenements_dict[row["Identifiant"]] = row["Code export paie"]
    
    # Dictionnaire pour les variables
    export_paie_variables_dict = {}
    if not export_paie_variables.empty:
        for _, row in export_paie_variables.iterrows():
            export_paie_variables_dict[row["Identifiant"]] = row["Code export paie"]
    
    return export_paie_evenements_dict, export_paie_variables_dict


def get_contrats_mapping(sheets):
    """Crée un dictionnaire pour mapper les identifiants aux contrats."""
    contrats_evenements = sheets.get("Contrats - événements", pd.DataFrame())
    contrats_variables = sheets.get("Contrats - variables de paie", pd.DataFrame())
    
    # Dictionnaire pour les événements : {identifiant_event: [(nom_contrat, regle_specifique)]}
    contrats_evenements_dict = {}
    if not contrats_evenements.empty:
        for _, row in contrats_evenements.iterrows():
            event_id = row["Identifiant"]
            if event_id not in contrats_evenements_dict:
                contrats_evenements_dict[event_id] = []
            contrats_evenements_dict[event_id].append({
                "nom_contrat": row["Nom du contrat"],
                "regle_specifique": row["Règle spécifique"]
            })
    
    # Dictionnaire pour les variables : {identifiant_variable: [(nom_contrat, regle_calcul)]}
    contrats_variables_dict = {}
    if not contrats_variables.empty:
        for _, row in contrats_variables.iterrows():
            var_id = row["Identifiant"]
            if var_id not in contrats_variables_dict:
                contrats_variables_dict[var_id] = []
            contrats_variables_dict[var_id].append({
                "nom_contrat": row["Nom du contrat"],
                "regle_calcul": row["Règle de calcul"]
            })
    
    return contrats_evenements_dict, contrats_variables_dict


def get_maj_mapping(sheets):
    """Crée un dictionnaire pour mapper les identifiants aux permissions client (MàJ)."""
    maj_evenements = sheets.get("MàJ compte événements", pd.DataFrame())
    maj_variables = sheets.get("MàJ compte variables de paie", pd.DataFrame())
    
    # Dictionnaire pour les événements : {identifiant: {peut_initialiser: bool, peut_mettre_a_jour: bool}}
    maj_evenements_dict = {}
    if not maj_evenements.empty:
        for _, row in maj_evenements.iterrows():
            event_id = row["Identifiant technique"]
            maj_evenements_dict[event_id] = {
                "peut_initialiser": row["Peut être initialisé par le client"],
                "peut_mettre_a_jour": row["Peut être mis à jour par le client"]
            }
    
    # Dictionnaire pour les variables : {identifiant: {peut_initialiser: bool, peut_mettre_a_jour: bool}}
    maj_variables_dict = {}
    if not maj_variables.empty:
        for _, row in maj_variables.iterrows():
            var_id = row["Identifiant technique"]
            maj_variables_dict[var_id] = {
                "peut_initialiser": row["Peut être initialisé par le client"],
                "peut_mettre_a_jour": row["Peut être mis à jour par le client"]
            }
    
    return maj_evenements_dict, maj_variables_dict


def get_notes_mapping(sheets):
    """Crée un dictionnaire pour mapper les noms aux notes."""
    notes_df = sheets.get("Liste des notes", pd.DataFrame())
    notes_dict = {}
    if not notes_df.empty:
        for _, row in notes_df.iterrows():
            notes_dict[row["Nom"]] = row["Note"]
    return notes_dict


# ============================================================================
# FONCTIONS POUR GÉNÉRER LE DOCUMENT WORD
# ============================================================================

def add_custom_styles(doc):
    """Ajoute des styles personnalisés au document Word."""
    styles = doc.styles
    
    # Style pour les titres de thème (niveau 1)
    theme_title_style = styles.add_style("ThemeTitle", WD_STYLE_TYPE.PARAGRAPH)
    theme_title_style.font.name = "Calibri"
    theme_title_style.font.size = Pt(16)
    theme_title_style.font.bold = True
    theme_title_style.font.color.rgb = PRIMARY_COLOR
    theme_title_style.paragraph_format.space_before = Pt(18)
    theme_title_style.paragraph_format.space_after = Pt(12)
    
    # Style pour les sous-sections (niveau 2)
    subsection_style = styles.add_style("SubsectionTitle", WD_STYLE_TYPE.PARAGRAPH)
    subsection_style.font.name = "Calibri"
    subsection_style.font.size = Pt(14)
    subsection_style.font.bold = True
    subsection_style.font.color.rgb = SECONDARY_COLOR
    subsection_style.paragraph_format.space_before = Pt(12)
    subsection_style.paragraph_format.space_after = Pt(6)
    
    # Style pour les listes à puces
    bullet_style = styles.add_style("BulletText", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(11)
    bullet_style.paragraph_format.left_indent = Pt(20)
    bullet_style.paragraph_format.space_before = Pt(2)
    bullet_style.paragraph_format.space_after = Pt(2)
    
    return doc


def add_theme_title(doc, title):
    """Ajoute un titre de thème (niveau 1)."""
    doc.add_paragraph(title, style="ThemeTitle")


def add_subsection_title(doc, title):
    """Ajoute un titre de sous-section (niveau 2)."""
    doc.add_paragraph(title, style="SubsectionTitle")


def add_bullet_point(doc, label, value):
    """Ajoute un point à une liste à puces."""
    if pd.notna(value) and value != "":
        doc.add_paragraph(f"{label} : {value}", style="BulletText")


def add_contrats_associes(doc, contrats_list):
    """Ajoute la liste des contrats associés."""
    if not contrats_list:
        return
    
    # Supprimer les doublons (basé sur nom_contrat + regle_specifique)
    contrats_unique = []
    seen = set()
    for contrat in contrats_list:
        nom_contrat = contrat.get("nom_contrat", "Inconnu")
        regle_specifique = contrat.get("regle_specifique", "Règle de base")
        key = (nom_contrat, regle_specifique)
        if key not in seen:
            seen.add(key)
            contrats_unique.append(contrat)
    
    if not contrats_unique:
        return
    
    doc.add_paragraph("Contrats associés :", style="BulletText")
    for contrat in contrats_unique:
        nom_contrat = contrat.get("nom_contrat", "Inconnu")
        regle_specifique = contrat.get("regle_specifique", "Règle de base")
        
        # Nettoyer la règle spécifique (enlever les sauts de ligne et espaces multiples)
        if pd.notna(regle_specifique):
            regle_specifique = str(regle_specifique).replace("\n", ", ").strip()
        
        if regle_specifique and regle_specifique != "nan" and regle_specifique != "Règle de base":
            doc.add_paragraph(f"  - {nom_contrat} : {regle_specifique}", style="BulletText")
        else:
            doc.add_paragraph(f"  - {nom_contrat} : Règle de base", style="BulletText")


def add_example_calcul(doc, nom, mode):
    """Ajoute un exemple de calcul basé sur le nom et le mode."""
    # Dictionnaire des exemples de calcul
    exemples = {
        "Congés Payés": "1 jour de CP = 7h",
        "Conges Payes": "1 jour de CP = 7h",
        "Récupération": "1h de récupération = 1h de temps libre",
        "Recuperation": "1h de récupération = 1h de temps libre",
        "Maladie": "1 jour de maladie = 0h (non rémunéré)",
        "Accident de Travail": "1 jour d'accident = 0h (selon convention)",
        "Conge Maternite": "1 jour de congé maternité = 0h (indemnisé par la Sécu)",
        "Conge Paternite": "1 jour de congé paternité = 0h (indemnisé par la Sécu)",
        "Conge Naissance": "1 jour de congé naissance = 0h (indemnisé)",
        "Conge Parental": "1 jour de congé parental = 0h (non rémunéré)",
        "Evenements Familiaux": "1 jour d'événement familial = 0h (selon convention)",
        "Enfant Malade": "1 jour d'enfant malade = 0h (selon convention)",
        "Accident de trajet": "1 jour d'accident de trajet = 0h (selon convention)",
        "Mi-temps therapeutique": "1 jour de mi-temps thérapeutique = 4h (50% du temps)",
        "Formation": "1h de formation = 1h (rémunérée)",
        "Absence Autorisee non payees": "1 jour d'absence autorisée = 0h (non rémunéré)",
        "Absence Non Autorisee": "1 jour d'absence non autorisée = 0h (sanction possible)",
        "Conge sans Solde": "1 jour de congé sans solde = 0h (non rémunéré)",
        "Presence automatique": "1 jour de présence automatique = 7h (selon planning)",
        "Mise a pied": "1 jour de mise à pied = 0h (non rémunéré)",
        "Heures de delegation CSE": "1h de délégation CSE = 1h (rémunérée)",
        "Heures Normales": "1h normale = 1h",
        "Heures Supplementaires 25%": "1h de HS25 = 1.25h de compensation",
        "Heures Supplementaires 50%": "1h de HS50 = 1.5h de compensation",
        "Heures Complementaires 10%": "1h de HC10% = 1.1h de compensation",
        "Heures de Nuit": "1h de nuit = 1.3h de compensation (selon convention)",
        "Heures de Dimanche": "1h de dimanche = 1.5h de compensation (selon convention)",
        "Heures de Ferie travaillees": "1h de férié travaillé = 1.5h de compensation",
        "Indemnite Casse Croute": "1 indemnité = 5€ (montant fixe)",
        "PRIME SAMEDI": "1 prime samedi = 10€ (montant fixe)",
        "Prime remplacement Admin": "1 prime = 15€ (montant fixe)",
        "Prime Astreinte maintenance": "1 prime astreinte = 20€ (montant fixe)",
        "COMPTEUR HEURE": "1h = 1h (compteur standard)",
    }
    
    # Nettoyer le nom (enlever les codes entre parenthèses et les espaces multiples)
    nom_clean = nom.split(" (")[0].strip()
    
    # Chercher dans le dictionnaire
    if nom_clean in exemples:
        doc.add_paragraph(f"Exemple de calcul : {exemples[nom_clean]}", style="BulletText")
    elif mode == "Avoir":
        doc.add_paragraph(f"Exemple de calcul : 1 jour = 7h (à adapter selon convention)", style="BulletText")
    elif mode == "Compensation":
        # Extraire le pourcentage si présent dans le nom
        if "%" in nom_clean:
            pourcentage = nom_clean.split("%")[0].split()[-1]
            if pourcentage.isdigit():
                doc.add_paragraph(f"Exemple de calcul : 1h = {1 + float(pourcentage)/100}h de compensation", style="BulletText")
            else:
                doc.add_paragraph(f"Exemple de calcul : 1h = Xh de compensation (à adapter)", style="BulletText")
        else:
            doc.add_paragraph(f"Exemple de calcul : 1h = 1h de compensation (à adapter)", style="BulletText")
    else:
        doc.add_paragraph(f"Exemple de calcul : 1 unité = 1h (à adapter)", style="BulletText")


# ============================================================================
# FONCTION PRINCIPALE POUR GÉNÉRER LE DOCUMENT
# ============================================================================

def generate_word_document(sheets):
    """Génère le document Word complet."""
    doc = Document()
    doc = add_custom_styles(doc)
    
    # Ajouter un titre principal
    doc.add_heading("Configuration Zeus - Données de l'entreprise", level=0)
    doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}")
    doc.add_paragraph()
    
    # Ajouter un résumé des thèmes
    doc.add_heading("Résumé des thèmes", level=1)
    theme_summary = (
        "Ce document décrit la configuration complète des données de l'entreprise dans Zeus. "
        "Il est organisé en thèmes correspondant aux onglets du fichier Excel source.\n"
        "\n"
        "- **Événements journaliers** : Liste des événements (CP, Maladie, Récupération, etc.) avec leurs règles, modes, et contrats associés.\n"
        "- **Liste des compteurs spéciaux** : Compteurs (Heures Normales, HS25%, etc.) avec leurs règles de calcul et permissions.\n"
        "- **Droits du logiciel** : Rôles (Manager, Responsable) et leurs permissions dans Zeus.\n"
        "- **Liste des notes** : Notes internes et remarques spécifiques.\n"
        "\n"
        "Chaque élément est décrit avec ses caractéristiques, son mode (Avoir/Compensation/Normal), "
        "et des exemples de calcul."
    )
    doc.add_paragraph(theme_summary)
    doc.add_paragraph()
    
    # Ajouter une table des matières (manuelle, car python-docx ne supporte pas les TOC automatiques)
    doc.add_heading("Table des matières", level=1)
    toc = (
        "1. Événements journaliers ............................................ 4\n"
        "2. Liste des compteurs spéciaux ..................................... X\n"
        "3. Droits du logiciel ................................................ X\n"
        "4. Liste des notes .................................................... X"
    )
    doc.add_paragraph(toc)
    doc.add_paragraph()
    doc.add_page_break()
    
    # Charger les mappings
    export_paie_evenements_dict, export_paie_variables_dict = get_export_paie_mapping(sheets)
    contrats_evenements_dict, contrats_variables_dict = get_contrats_mapping(sheets)
    maj_evenements_dict, maj_variables_dict = get_maj_mapping(sheets)
    notes_dict = get_notes_mapping(sheets)
    
    # Traiter chaque onglet (thème)
    for sheet_name in sheets.keys():
        df = sheets[sheet_name]
        
        # Ignorer les onglets vides, Export paie, Contrats, ou MàJ (car déjà inclus dans les autres sections)
        if df.empty or sheet_name.startswith("Export paie") or sheet_name.startswith("Contrats") or sheet_name.startswith("MàJ"):
            continue
        
        # Ajouter le titre du thème
        add_theme_title(doc, sheet_name)
        
        # Traiter chaque ligne comme une sous-section
        for _, row in df.iterrows():
            # Déterminer le nom de la sous-section
            if "Nom" in df.columns:
                nom = row["Nom"]
            elif "Nom court" in df.columns:
                nom = row["Nom court"]
            elif "Identifiant" in df.columns:
                nom = f"ID {row['Identifiant']}"
            else:
                nom = f"Élément {_ + 1}"
            
            # Ajouter le titre de la sous-section
            add_subsection_title(doc, nom)
            
            # Déterminer le mode
            mode = determine_mode(
                row, sheet_name, 
                contrats_evenements_df=sheets.get("Contrats - événements"),
                contrats_variables_df=sheets.get("Contrats - variables de paie"),
                compteurs_speciaux_df=sheets.get("Liste des compteurs spéciaux")
            )
            
            # Ajouter les champs spécifiques selon le thème
            if sheet_name == "Événements journaliers":
                add_bullet_point(doc, "Identifiant", row.get("Identifiant"))
                add_bullet_point(doc, "Nom court", row.get("Nom court"))
                add_bullet_point(doc, "Code", row.get("Code"))
                
                # Code export paie
                event_id = row.get("Identifiant")
                code_export = export_paie_evenements_dict.get(event_id)
                add_bullet_point(doc, "Code export paie", code_export)
                
                if mode:
                    add_bullet_point(doc, "Mode", mode)
                
                add_bullet_point(doc, "Règle d'acquisition", row.get("Règle d'acquisition"))
                add_bullet_point(doc, "Règle de prise", row.get("Règle de prise"))
                add_bullet_point(doc, "Cas particulier", row.get("Cas particulier"))
                add_bullet_point(doc, "Calendaire", row.get("Calendaire"))
                # Permissions client (de l'onglet Événements journaliers)
                add_bullet_point(doc, "Le client peut initialiser", row.get("Le client peut initialiser"))
                add_bullet_point(doc, "Le client peut mettre à jour", row.get("Le client peut mettre à jour"))
                
                # Permissions client (de l'onglet MàJ compte événements)
                maj_info = maj_evenements_dict.get(event_id)
                if maj_info:
                    add_bullet_point(doc, "Peut être initialisé par le client (MàJ)", maj_info.get("peut_initialiser"))
                    add_bullet_point(doc, "Peut être mis à jour par le client (MàJ)", maj_info.get("peut_mettre_a_jour"))
                
                # Contrats associés
                contrats_list = contrats_evenements_dict.get(event_id, [])
                add_contrats_associes(doc, contrats_list)
                
                # Compteur à déduire
                if "Compteur à déduire" in df.columns:
                    add_bullet_point(doc, "Compteur à déduire", row.get("Compteur à déduire"))
                
                # Note : vérifier si le nom ou le code correspond à une note
                note = None
                if nom in notes_dict:
                    note = notes_dict[nom]
                elif row.get("Code") in notes_dict:
                    note = notes_dict[row.get("Code")]
                elif row.get("Nom court") in notes_dict:
                    note = notes_dict[row.get("Nom court")]
                
                if note:
                    add_bullet_point(doc, "Note", note)
                
                # Vérifier si la note "CONTRAT 39H" s'applique à cet événement
                if "39H" in str(row.get("Code", "")) or "39H" in str(row.get("Nom", "")):
                    contrat_39h_note = notes_dict.get("CONTRAT 39H")
                    if contrat_39h_note:
                        add_bullet_point(doc, "Note (CONTRAT 39H)", contrat_39h_note)
                
                # Exemple de calcul
                add_example_calcul(doc, nom, mode)
                
            elif sheet_name == "Liste des compteurs spéciaux":
                add_bullet_point(doc, "Identifiant", row.get("Identifiant"))
                add_bullet_point(doc, "Nom court", row.get("Nom court"))
                add_bullet_point(doc, "Code", row.get("Code"))
                
                # Code export paie
                compteur_id = row.get("Identifiant")
                code_export = export_paie_variables_dict.get(compteur_id)
                add_bullet_point(doc, "Code export paie", code_export)
                
                if mode:
                    add_bullet_point(doc, "Mode", mode)
                
                add_bullet_point(doc, "Règle de calcul", row.get("Règle de calcul"))
                add_bullet_point(doc, "Remarque", row.get("Remarque"))
                add_bullet_point(doc, "Le client peut initialiser", row.get("Le client peut initialiser"))
                add_bullet_point(doc, "Le client peut mettre à jour", row.get("Le client peut mettre à jour"))
                
                # Permissions client (de l'onglet MàJ compte variables de paie)
                maj_info = maj_variables_dict.get(compteur_id)
                if maj_info:
                    add_bullet_point(doc, "Peut être initialisé par le client (MàJ)", maj_info.get("peut_initialiser"))
                    add_bullet_point(doc, "Peut être mis à jour par le client (MàJ)", maj_info.get("peut_mettre_a_jour"))
                
                # Contrats associés (via Contrats - variables de paie)
                contrats_list = contrats_variables_dict.get(compteur_id, [])
                if contrats_list:
                    add_contrats_associes(doc, contrats_list)
                
                # Exemple de calcul
                add_example_calcul(doc, nom, mode)
                
            elif sheet_name == "Contrats - événements":
                add_bullet_point(doc, "Identifiant", row.get("Identifiant"))
                add_bullet_point(doc, "Nom du contrat", row.get("Nom du contrat"))
                add_bullet_point(doc, "Code", row.get("Code"))
                add_bullet_point(doc, "Nom", row.get("Nom"))
                add_bullet_point(doc, "Code.1", row.get("Code.1"))
                
                if mode:
                    add_bullet_point(doc, "Mode", mode)
                
                add_bullet_point(doc, "Règle spécifique", row.get("Règle spécifique"))
                add_bullet_point(doc, "Règle d'acquisition", row.get("Règle d'acquisition"))
                add_bullet_point(doc, "Règle de prise", row.get("Règle de prise"))
                add_bullet_point(doc, "Cas particulier", row.get("Cas particulier"))
                add_bullet_point(doc, "Compteur à déduire", row.get("Compteur à déduire"))
                
            elif sheet_name == "Contrats - variables de paie":
                add_bullet_point(doc, "Identifiant", row.get("Identifiant"))
                add_bullet_point(doc, "Nom du contrat", row.get("Nom du contrat"))
                add_bullet_point(doc, "Code", row.get("Code"))
                add_bullet_point(doc, "Nom", row.get("Nom"))
                add_bullet_point(doc, "Code.1", row.get("Code.1"))
                
                if mode:
                    add_bullet_point(doc, "Mode", mode)
                
                add_bullet_point(doc, "Règle spécifique", row.get("Règle spécifique"))
                add_bullet_point(doc, "Règle de calcul", row.get("Règle de calcul"))
                
            elif sheet_name == "MàJ compte événements":
                add_bullet_point(doc, "Identifiant technique", row.get("Identifiant technique"))
                add_bullet_point(doc, "Nom", row.get("Nom"))
                add_bullet_point(doc, "Peut être initialisé par le client", row.get("Peut être initialisé par le client"))
                add_bullet_point(doc, "Peut être mis à jour par le client", row.get("Peut être mis à jour par le client"))
                
                # Lier à l'événement correspondant
                event_id = row.get("Identifiant technique")
                if event_id in sheets.get("Événements journaliers", pd.DataFrame())["Identifiant"].values:
                    event_name = sheets["Événements journaliers"][
                        sheets["Événements journaliers"]["Identifiant"] == event_id
                    ]["Nom"].iloc[0]
                    add_bullet_point(doc, "Lié à l'événement", event_name)
                
            elif sheet_name == "MàJ compte variables de paie":
                add_bullet_point(doc, "Identifiant technique", row.get("Identifiant technique"))
                add_bullet_point(doc, "Nom", row.get("Nom"))
                add_bullet_point(doc, "Peut être initialisé par le client", row.get("Peut être initialisé par le client"))
                add_bullet_point(doc, "Peut être mis à jour par le client", row.get("Peut être mis à jour par le client"))
                
                # Lier à la variable correspondant
                var_id = row.get("Identifiant technique")
                if var_id in sheets.get("Liste des compteurs spéciaux", pd.DataFrame())["Identifiant"].values:
                    var_name = sheets["Liste des compteurs spéciaux"][
                        sheets["Liste des compteurs spéciaux"]["Identifiant"] == var_id
                    ]["Nom"].iloc[0]
                    add_bullet_point(doc, "Lié à la variable", var_name)
                
            elif sheet_name == "Droits du logiciel":
                add_bullet_point(doc, "Rôle", row.get("Rôle"))
                add_bullet_point(doc, "Nom", row.get("Nom"))
                add_bullet_point(doc, "Commentaire", row.get("Commentaire"))
                add_bullet_point(doc, "Peut pointer pour son équipe", row.get("Peut pointer pour son équipe"))
                add_bullet_point(doc, "Peut demander un pointage oublié", row.get("Peut demander un pointage oublié"))
                add_bullet_point(doc, "Événements que je peux poser", row.get("Événements que je peux poser"))
                add_bullet_point(doc, "Événements que je peux voir", row.get("Événements que je peux voir"))
                add_bullet_point(doc, "Événements que je peux demander", row.get("Événements que je peux demander"))
                
            elif sheet_name == "Liste des notes":
                add_bullet_point(doc, "Nom", row.get("Nom"))
                add_bullet_point(doc, "Note", row.get("Note"))
            
            # Ajouter un espace entre les sous-sections
            doc.add_paragraph()
        
        # Ajouter un espace entre les thèmes
        doc.add_paragraph()
    
    return doc


# ============================================================================
# EXÉCUTION PRINCIPALE
# ============================================================================

if __name__ == "__main__":
    print("Chargement du fichier Excel...")
    sheets = load_excel_data(EXCEL_FILE)
    
    print("Génération du document Word...")
    doc = generate_word_document(sheets)
    
    print(f"Sauvegarde du document : {OUTPUT_WORD_FILE}")
    doc.save(OUTPUT_WORD_FILE)
    
    print(f"✅ Document généré avec succès : {OUTPUT_WORD_FILE}")
